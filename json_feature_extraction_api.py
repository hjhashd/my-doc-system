# -*- coding: utf-8 -*-
import json
import re
import shutil
from pathlib import Path
import pandas as pd
from urllib.parse import quote

from docx import Document
from PIL import Image
import os

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor

# ====== 可调阈值 ======
LARGE_TEXT_MIN_W = 500   # 把 text 纠正为 image 的最小宽
LARGE_TEXT_MIN_H = 400   # 把 text 纠正为 image 的最小高
CONTAIN_PAD = 2          # image 包含判断的像素容差
IOU_DROP = 0.9           # image IoU 超过该阈值则保留更大的一张


# ====== 字体：宋体设置 ======
def set_doc_font_simsun(doc: Document, font_name: str = "SimSun"):
    """
    将文档默认样式（Normal）与超链接样式（Hyperlink）设置为宋体。
    同时设置 eastAsia 族，确保中文显示为宋体。
    """
    # Normal
    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        # eastAsia
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        # 英文/西文
        normal.font._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        normal.font._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass

    # Hyperlink（有些模板可能没有该样式，做 try 保护）
    try:
        hyperlink = doc.styles["Hyperlink"]
        hyperlink.font.name = font_name
        hyperlink._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        hyperlink._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        hyperlink._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass


# ====== 超链接（使用相对路径） ======
def add_hyperlink(paragraph, text: str, target_ref: str, font_name: str = "SimSun"):
    """
    向段落 paragraph 插入一个超链接，文本显示为 text，链接目标使用 target_ref（相对路径字符串）。
    注意：这里不再使用 file:/// 绝对 URI，而是直接把相对路径写进外部关系中，
    Word 会以 .docx 所在目录为基准解析它。
    """
    part = paragraph.part
    # target_ref 例如 "img/XA_certificate_0_layout_det_res_1.png"
    r_id = part.relate_to(target_ref, RT.HYPERLINK, is_external=True)

    # <w:hyperlink r:id="...">
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # <w:r><w:rPr>...</w:rPr><w:t>text</w:t></w:r>
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 使用 Hyperlink 样式（蓝色+下划线）
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # 显式设置字体族为宋体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


# ====== 基础工具 ======
def load_json(json_path: Path):
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_page_tag(data: dict):
    input_path = data["input_path"]
    page_index = data["page_index"]
    pdf_name = Path(input_path).name
    base_name = pdf_name.rsplit(".", 1)[0]
    page_tag = f"{base_name}_{page_index}"
    return base_name, page_index, page_tag


def collect_header_bands(parsing_res_list):
    bands = []
    for blk in parsing_res_list:
        if blk.get("block_label") == "header":
            x1, y1, x2, y2 = blk["block_bbox"]
            bands.append((float(y1), float(y2)))
    return bands


def _y_overlap_ratio(a, b):
    a1, a2 = min(a), max(a)
    b1, b2 = min(b), max(b)
    inter = max(0.0, min(a2, b2) - max(a1, b1))
    denom = max(1e-6, min(a2 - a1, b2 - b1))
    return inter / denom


def is_in_header_band(block, header_bands, overlap_ratio_thresh=0.5):
    if not header_bands:
        return False
    _, y1, _, y2 = block["block_bbox"]
    for hy1, hy2 in header_bands:
        if _y_overlap_ratio((y1, y2), (hy1, hy2)) >= overlap_ratio_thresh:
            return True
    return False


def ensure_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)


def crop_image_region(layout_png_path: Path, bbox, out_path: Path):
    if not layout_png_path.exists():
        print(f"[WARN] layout image not found: {layout_png_path}")
        return False

    img = Image.open(layout_png_path)
    x1, y1, x2, y2 = bbox
    x1 = max(0, int(x1)); y1 = max(0, int(y1))
    x2 = min(img.width, int(x2)); y2 = min(img.height, int(y2))
    if x2 <= x1 or y2 <= y1:
        print(f"[WARN] invalid bbox, skip crop: {bbox}")
        return False

    cropped = img.crop((x1, y1, x2, y2))
    cropped.save(out_path)
    return True


# ====== 文档与路径处理 ======
def copy_table_file(asset_base_dir: Path, page_tag: str, table_idx: int, word_doc: Document, 
                    table_dir: Path, link_base_dir: Path, url_prefix: str = ""):
    """
    将 Excel 复制到输出目录，生成美化链接 + 隐形文件名锚点
    """
    src_name = f"{page_tag}_table_{table_idx}.xlsx"
    src_path = asset_base_dir / src_name

    ensure_dir(table_dir)
    dst_path = table_dir / src_name

    # 1. 物理复制
    if src_path.exists():
        shutil.copy2(src_path, dst_path)
    else:
        print(f"[WARN] table xlsx not found: {src_path}")
        open(dst_path, "a").close()

    # 2. 生成 Web 跳转链接
    if url_prefix:
        parts = url_prefix.strip("/").split("/")
        agent_user_id = parts[1] if len(parts) >= 3 else ""
        task_id = parts[2] if len(parts) >= 3 else ""
        prefix = url_prefix.rstrip("/")
        file_web_path = f"{prefix}/table/{src_name}"
        encoded_url = quote(file_web_path)
        encoded_name = quote(src_name)
        rel_target = f"/excel-editor?docUrl={encoded_url}&docName={encoded_name}&agentUserId={agent_user_id}&taskId={task_id}"
    else:
        rel = os.path.relpath(dst_path, start=link_base_dir)
        rel_target = rel.replace("\\", "/")   

    # 3. 写入 Word (关键修改)
    p = word_doc.add_paragraph()
    
    # A. 【给人看】美化后的文字
    link_text = f"� 点击编辑表格 ({src_name})"
    add_hyperlink(p, link_text, rel_target)

    # B. 【给程序看】添加隐形签名 {{#T#:文件名}}
    hidden_signature = f"{{{{#T#:{src_name}}}}}"  
    run = p.add_run(hidden_signature)
    
    # ====== 修改开始：真正的视觉隐藏 (白色 + 极小字号) ====== 
    run.font.size = Pt(0.5)                       # 字号设为 0.5 磅 (肉眼几乎不可见) 
    run.font.color.rgb = RGBColor(255, 255, 255)  # 颜色设为白色 (背景也是白的话就看不见了) 
    # ====== 修改结束 ====== 
    
    # 设置隐藏属性 (保留这个属性，双重保险)
    rPr = run._element.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)


def sort_blocks_reading_order(parsing_res_list, y_tol=10):
    """
    按阅读顺序排序：先 y（分桶），再 x 升序
    """
    def key_fn(blk):
        x1, y1, x2, y2 = blk.get("block_bbox", [0, 0, 0, 0])
        row = round(float(y1) / float(y_tol))
        return (row, float(x1))
    return sorted(parsing_res_list, key=key_fn)


def is_other_text_attached(title_blk, cand_blk, y_tol=10):
    """
    判断 other_text 是否应附加到 paragraph_title：
      - y 高度重叠 > 50%
      - 或者 中线距离 <= 2 * y_tol
    """
    _, ty1, _, ty2 = title_blk["block_bbox"]
    _, cy1, _, cy2 = cand_blk["block_bbox"]
    if _y_overlap_ratio((ty1, ty2), (cy1, cy2)) >= 0.5:
        return True
    if abs((ty1 + ty2) / 2.0 - (cy1 + cy2) / 2.0) <= 2 * y_tol:
        return True
    return False


# ====== bbox 工具函数 ======
def _bbox_area(b):
    x1, y1, x2, y2 = map(float, b)
    return max(0.0, x2 - x1) * max(0.0, y2 - y1)


def _bbox_contains(outer, inner, pad=0.0):
    ox1, oy1, ox2, oy2 = map(float, outer)
    ix1, iy1, ix2, iy2 = map(float, inner)
    return (ix1 >= ox1 - pad and iy1 >= oy1 - pad and
            ix2 <= ox2 + pad and iy2 <= oy2 + pad)


def _bbox_iou(a, b):
    ax1, ay1, ax2, ay2 = map(float, a)
    bx1, by1, bx2, by2 = map(float, b)
    ix1, iy1 = max(ax1, bx1), max(ay1, by1)
    ix2, iy2 = min(ax2, bx2), min(ay2, by2)
    iw, ih = max(0.0, ix2 - ix1), max(0.0, iy2 - iy1)
    inter = iw * ih
    aa = max(0.0, ax2 - ax1) * max(0.0, ay2 - ay1)
    ba = max(0.0, bx2 - bx1) * max(0.0, by2 - by1)
    union = aa + ba - inter + 1e-6
    return inter / union


# ====== 规则修正：text -> image（大块） ======
def coerce_large_text_to_image(parsing_res_list, min_w=LARGE_TEXT_MIN_W, min_h=LARGE_TEXT_MIN_H):
    new_list = []
    for blk in parsing_res_list:
        label = blk.get("block_label")
        if label == "text":
            x1, y1, x2, y2 = blk.get("block_bbox", [0, 0, 0, 0])
            w, h = float(x2 - x1), float(y2 - y1)
            if w >= min_w and h >= min_h:
                blk = dict(blk)
                blk["block_label"] = "image"
        new_list.append(blk)
    return new_list


# ====== 仅保留更大 image：去包含/高 IoU ======
def drop_nested_images(parsing_res_list, contain_pad=CONTAIN_PAD, iou_drop=IOU_DROP):
    images, others = [], []
    for blk in parsing_res_list:
        (images if blk.get("block_label") == "image" else others).append(blk)

    images_sorted = sorted(images, key=lambda b: _bbox_area(b["block_bbox"]), reverse=True)
    kept = []
    for b in images_sorted:
        bb = b["block_bbox"]
        drop = False
        for k in kept:
            kb = k["block_bbox"]
            if _bbox_contains(kb, bb, pad=contain_pad):
                drop = True
                break
            if _bbox_iou(kb, bb) >= iou_drop:
                drop = True
                break
        if not drop:
            kept.append(b)

    kept_set = set(id(x) for x in kept)
    images_kept_in_original_order = [blk for blk in images if id(blk) in kept_set]
    return images_kept_in_original_order + others


# ====== 处理单个 JSON 并追加到 Word ======
def process_one_json_and_append(doc: Document, json_path: Path, img_dir: Path, table_dir: Path, link_base_dir: Path, url_prefix: str = ""):
    data = load_json(json_path)
    base_name, page_index, page_tag = get_page_tag(data)

    print(f"processing JSON: {json_path}")
    print(f"base_name = {base_name}, page_index = {page_index}, page_tag = {page_tag}")

    parsing_res_list = data.get("parsing_res_list", [])

    # 规则修正 + 排序
    parsing_res_list = coerce_large_text_to_image(parsing_res_list)
    parsing_res_list = drop_nested_images(parsing_res_list)
    parsing_res_list = sort_blocks_reading_order(parsing_res_list, y_tol=10)

    header_bands = collect_header_bands(parsing_res_list)

    # 以 JSON 文件所在目录为基准寻找资源
    asset_base_dir = json_path.parent
    layout_png_path = asset_base_dir / f"{page_tag}_layout_det_res.png"

    image_idx = 1
    table_idx = 1
    i = 0
    n = len(parsing_res_list)

    ensure_dir(img_dir)
    ensure_dir(table_dir)

    while i < n:
        blk = parsing_res_list[i]
        label = blk.get("block_label")
        content = blk.get("block_content", "") or ""

        if label == "header" or is_in_header_band(blk, header_bands, overlap_ratio_thresh=0.5) or label == "seal":
            i += 1
            continue

        if label == "paragraph_title":
            merged_text = re.sub(r"\s+", "", content or "")
            j = i + 1
            while j < n:
                next_blk = parsing_res_list[j]
                if (next_blk.get("block_label") == "other_text"
                    and not is_in_header_band(next_blk, header_bands, overlap_ratio_thresh=0.5)
                    and is_other_text_attached(blk, next_blk, y_tol=10)):
                    merged_text += re.sub(r"\s+", "", next_blk.get("block_content", "") or "")
                    j += 1
                else:
                    break
            doc.add_paragraph(merged_text)
            i = j
            continue

        if label == "text":
            merged_text = content
            j = i + 1
            while j < n:
                next_blk = parsing_res_list[j]
                if (next_blk.get("block_label") == "other_text"
                    and not is_in_header_band(next_blk, header_bands, overlap_ratio_thresh=0.5)
                    and is_other_text_attached(blk, next_blk, y_tol=10)):
                    merged_text += next_blk.get("block_content", "") or ""
                    j += 1
                else:
                    break
            doc.add_paragraph(merged_text)
            i = j
            continue

        if label == "image":
            out_name = f"{page_tag}_layout_det_res_{image_idx}.png"
            out_path = img_dir / out_name
            if crop_image_region(layout_png_path, blk["block_bbox"], out_path):
                # 1. 计算链接目标
                if url_prefix:
                    prefix = url_prefix.rstrip("/")
                    rel_target = f"{prefix}/img/{out_name}"
                else:
                    rel = os.path.relpath(out_path, start=link_base_dir)
                    rel_target = rel.replace("\\", "/")
                
                p = doc.add_paragraph()
                
                # A. 【给人看】美化后的文字
                visible_text = f"🖼️ 点击查看图片 ({out_name})"
                add_hyperlink(p, visible_text, rel_target)
                
                # B. 【给程序看】添加隐形签名 {{#I#:文件名}}
                hidden_signature = f"{{{{#I#:{out_name}}}}}"
                run = p.add_run(hidden_signature)
                
                # ====== 修改开始：真正的视觉隐藏 (白色 + 极小字号) ====== 
                run.font.size = Pt(0.5)                       # 字号设为 0.5 磅 
                run.font.color.rgb = RGBColor(255, 255, 255)  # 颜色设为白色 
                # ====== 修改结束 ====== 
                
                # 设置隐藏属性
                rPr = run._element.get_or_add_rPr()
                vanish = OxmlElement('w:vanish')
                rPr.append(vanish)

                image_idx += 1
            else:
                if content.strip():
                    doc.add_paragraph(content.strip())
            i += 1
            continue

        if label == "table":
            # ====== [修改点] 传递 url_prefix 给表格函数 ====== 
            copy_table_file(asset_base_dir, page_tag, table_idx, doc, table_dir, 
                            link_base_dir=link_base_dir, url_prefix=url_prefix)
            # ================================================ 
            table_idx += 1
            i += 1
            continue

        if content.strip():
            doc.add_paragraph(content.strip())

        i += 1

    return base_name


def parse_file_sort_key(p: Path):
    """
    解析 *_res.json 的排序键：(base_name, index)
    若文件名匹配失败，回退读取 JSON 中的 page_index
    """
    m = re.match(r"^(?P<base>.+)_(?P<idx>\d+)_res\.json$", p.name)
    if m:
        base = m.group("base")
        idx = int(m.group("idx"))
        return (base, idx)
    try:
        data = load_json(p)
        base_name, page_index, _ = get_page_tag(data)
        return (base_name, int(page_index))
    except Exception:
        return (p.stem, 0)


# ====== 可复用生成函数 ======
def generate_word_from_jsons(json_dir: Path, img_dir: Path, table_dir: Path, out_docx_dir: Path, url_prefix: str = ""):
    """
    扫描 json_dir 下的 *_res.json，按页序合并写入一个 Word。
    图片裁剪保存到 img_dir，表格复制到 table_dir，最终 Word 保存在 out_docx_dir。
    如果提供了 url_prefix，超链接目标使用该前缀开头的绝对路径；否则使用相对路径。
    返回 (out_docx_path, True)。
    """
    if not json_dir.exists() or not json_dir.is_dir():
        raise FileNotFoundError(f"not a directory: {json_dir}")

    json_files = sorted(json_dir.glob("*_res.json"), key=parse_file_sort_key)
    if not json_files:
        raise FileNotFoundError(f"no *_res.json under {json_dir}")

    ensure_dir(img_dir)
    ensure_dir(table_dir)
    ensure_dir(out_docx_dir)

    doc = Document()
    # 统一设置为宋体
    set_doc_font_simsun(doc, "SimSun")

    base_names_seen = []
    for jp in json_files:
        base_name = process_one_json_and_append(doc, jp, img_dir=img_dir, table_dir=table_dir,
                                                link_base_dir=out_docx_dir, url_prefix=url_prefix)
        base_names_seen.append(base_name)

    uniq_bases = list(dict.fromkeys(base_names_seen))
    out_name = f"{uniq_bases[0]}_res.docx" if len(uniq_bases) == 1 else "merged_res.docx"
    out_docx = out_docx_dir / out_name

    doc.save(out_docx)
    print(f"saved Word: {out_docx.resolve()}")
    return out_docx, True


# output_base_dir = "./pdf_output/"
# output_pdf_id_path = os.path.join(output_base_dir, "2")

# pdf_out_path = output_pdf_id_path
# file_name = "XA_certificate"
# json_dir = Path(pdf_out_path) / file_name
# img_dir =  Path(pdf_out_path+"/img")
# table_dir = Path(pdf_out_path+"/table")
# out_docx_dir = Path(pdf_out_path+"/")

# print(json_dir)
# print(img_dir)
# print(table_dir)
# print(out_docx_dir)

# generate_word_from_jsons(json_dir, img_dir, table_dir, out_docx_dir)